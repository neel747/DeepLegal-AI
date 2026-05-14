import os
from docx import Document

def create_sample_docx(filename, content):
    doc = Document()
    for line in content:
        if line.startswith("Section"):
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)
    doc.save(filename)

contracts_dir = "data/contracts"
os.makedirs(contracts_dir, exist_ok=True)

# Sample NDA
nda_content = [
    "NON-DISCLOSURE AGREEMENT",
    "This Agreement is made as of May 14, 2026, between Client Corp and Vendor Inc.",
    "Section 1. Confidential Information",
    "Confidential Information shall include all data provided by the Disclosing Party...",
    "Section 2. Obligations",
    "The Receiving Party shall maintain the confidentiality of the information...",
    "Section 3. Termination",
    "This agreement expires in 2 years from the effective date."
]
create_sample_docx(os.path.join(contracts_dir, "sample_nda.docx"), nda_content)

# Sample MSA
msa_content = [
    "MASTER SERVICES AGREEMENT",
    "Section 1. Scope of Work",
    "Vendor shall provide services as described in subsequent SOWs...",
    "Section 2. Fees and Payment",
    "Client shall pay Vendor $10,000 per month...",
    "Section 3. Liability",
    "The total liability shall not exceed $1,000,000."
]
create_sample_docx(os.path.join(contracts_dir, "sample_msa.docx"), msa_content)

print("Sample contracts created in data/contracts/")
